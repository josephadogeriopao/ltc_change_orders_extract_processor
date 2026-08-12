"""
Barcode Generation Unit Test Suite
----------------------------------
Description: Builds a temporary mock template document, applies the inline
             barcode stamp execution routines, and verifies successful replacement.

File: tests/test_barcode_generator.py
"""

import os
import sys
import unittest
from docx import Document

# Safe absolute tracking path injection to import from src/ seamlessly
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "src")))
from src.utils.barcode_generator import insert_code128_barcode


class TestBarcodeGenerator(unittest.TestCase):

    def setUp(self):
        """Creates an isolated runtime directory and builds a dummy template."""
        self.test_dir = os.path.abspath(os.path.dirname(__file__))
        self.doc_path = os.path.join(self.test_dir, "mock_template.docx")
        self.output_path = os.path.join(self.test_dir, "mock_output.docx")
        self.parcel_ids = ["12345ABC", "67890XYZ"]

        # Build a temporary Word doc containing the placeholders
        doc = Document()
        doc.add_paragraph("This is an inline parcel tracking reference: BARCODETARGET")

        # Add a placeholder inside a table layout cell too
        table = doc.add_table(rows=1, cols=1)
        table.rows[0].cells[0].paragraphs[0].text = "Table layout index: BARCODETARGET"

        doc.save(self.doc_path)

    def tearDown(self):
        """Wipes temporary generated files off the workspace clean."""
        for path in [self.doc_path, self.output_path]:
            if os.path.exists(path):
                os.remove(path)

    def test_barcode_stamping_removes_placeholders(self):
        """Verifies placeholders are extracted and replaced successfully."""
        # Execute the barcode engine runner
        insert_code128_barcode(
            doc_path=self.doc_path,
            output_path=self.output_path,
            all_parcel_ids=self.parcel_ids,
            output_dir=self.test_dir
        )

        # Assert output file layer compiled successfully
        self.assertTrue(os.path.exists(self.output_path))

        # Re-parse output to confirm 'BARCODETARGET' text is gone
        output_doc = Document(self.output_path)

        # Inspect paragraphs scope
        for paragraph in output_doc.paragraphs:
            self.assertNotIn("BARCODETARGET", paragraph.text)

        # Inspect tables cell scope
        for table in output_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self.assertNotIn("BARCODETARGET", paragraph.text)


if __name__ == "__main__":
    unittest.main()
