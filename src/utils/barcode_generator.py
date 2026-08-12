"""
Job Configuration and Registry Blueprint
-----------------------------------------
Description: Defines centralized property configuration tracking types and uses
             an execution-aware resolver to safely bind document template assets.

Author: Joseph Adogeri
Version: 1.0.0
Since: 2026-08-03
File: job_config.py / constants.py
License: MIT
"""

import os
from docx import Document
from docx.shared import Inches
import barcode
from barcode.writer import ImageWriter


def insert_code128_barcode(doc_path, output_path, all_parcel_ids, output_dir):
    """Finds 'BARCODETARGET' text inside a unified master Word document and replaces

    each consecutive instance with its corresponding unique Code 128 barcode image.
    """
    if not all_parcel_ids:
        print(f"  ⚠️ Skipping barcode generation: Missing Parcel ID database.")
        return

    # Exact metric to inch conversion markers (1 inch = 25.4 mm)
    TARGET_HEIGHT_INCHES = 3.175 / 25.4
    TARGET_WIDTH_INCHES = 39.917 / 25.4

    code128 = barcode.get_barcode_class('code128')
    options = {'write_text': False, 'module_height': 15.0, 'module_width': 0.2}

    # Open the master merged Word document
    doc = Document(doc_path)
    placeholder = "BARCODETARGET"

    # Counter track to map sequential occurrences to our parcel ID list index array
    barcode_index = 0
    total_ids = len(all_parcel_ids)

    def process_runs_with_unique_barcode(paragraph_obj, current_index):
        """Helper to generate a specific barcode image on-the-fly and swap it."""
        nonlocal barcode_index

        # Guard clause if the document somehow contains more target strings than data rows
        if barcode_index >= total_ids:
            return

        target_pid = all_parcel_ids[barcode_index]

        # Generate the unique barcode image file
        temp_filename = os.path.join(output_dir, f"temp_bc_{barcode_index}_{os.getpid()}")
        barcode_instance = code128(target_pid, writer=ImageWriter())
        barcode_instance.save(temp_filename, options=options)
        full_image_path = f"{temp_filename}.png"

        for run in paragraph_obj.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, "")
                run.add_picture(
                    full_image_path,
                    width=Inches(TARGET_WIDTH_INCHES),
                    height=Inches(TARGET_HEIGHT_INCHES)
                )

        # Clean up the specific temp image instantly
        if os.path.exists(full_image_path):
            os.remove(full_image_path)

        print(f"  > Page Link [{barcode_index + 1}] Embedded barcode for ID: {target_pid}")
        barcode_index += 1

    # Loop through body paragraphs sequentially
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            process_runs_with_unique_barcode(paragraph, barcode_index)

    # Loop through table cells sequentially (in case placeholder sits in layout boxes)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        process_runs_with_unique_barcode(paragraph, barcode_index)

    # Safe-save wrapper to catch file-lock permission errors cleanly
    try:
        doc.save(output_path)
        print(f"  ✅ Global inline-replacement finished. Stamped {barcode_index} unique barcodes.")
    except PermissionError:
        print(f"  ❌ ERROR: Cannot save file! Please close '{output_path}' in Word and retry.")
