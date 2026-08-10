import os
from docx import Document
from docx.shared import Inches
import barcode
from barcode.writer import ImageWriter


def insert_code128_barcode(doc_path, output_path, parcel_id):
    # 1. Define exact metric to inch conversion markers (1 inch = 25.4 mm)
    TARGET_HEIGHT_INCHES = 3.175 / 25.4  # Exact 0.125 inches
    TARGET_WIDTH_INCHES = 39.917 / 25.4  # Exact 1.5715 inches

    # 2. Initialize the Code 128 generator with an Image Writer (generates a PNG)
    code128 = barcode.get_barcode_class('code128')

    # 3. Configure visual options (removes the default text beneath the bars)
    options = {
        'write_text': False,
        'module_height': 15.0,
        'module_width': 0.2
    }

    # 4. Generate and save the temporary barcode image
    barcode_instance = code128(parcel_id, writer=ImageWriter())
    temp_filename = "temp_barcode"
    barcode_instance.save(temp_filename, options=options)
    full_image_path = f"{temp_filename}.png"

    # 5. Open your Word document
    doc = Document(doc_path)
    placeholder = "BARCODETARGET"

    # Helper function to swap text inline without altering paragraph structural bounds
    def process_paragraph(p):
        if placeholder in p.text:
            # REMOVED: p.alignment setting to keep the layout left-anchored exactly where it starts
            for run in p.runs:
                if placeholder in run.text:
                    # Clear out the text placeholder
                    run.text = run.text.replace(placeholder, "")

                    # Insert the barcode picture directly at the precise left-aligned start spot
                    run.add_picture(
                        full_image_path,
                        width=Inches(TARGET_WIDTH_INCHES),
                        height=Inches(TARGET_HEIGHT_INCHES)
                    )
                    # REMOVED: label_run block that appended data text underneath the bars

    # 6. Process body paragraphs
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)

    # 7. Process table cells (in case BARCODETARGET is inside layout tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)

    # 8. Safe-save wrapper to catch file-lock permission errors cleanly
    try:
        doc.save(output_path)
        print(f"Successfully generated and inline-replaced '{placeholder}' with barcode for {parcel_id}!")
    except PermissionError:
        print(
            f"\n❌ ERROR: Cannot save file! Please close '{output_path}' in Microsoft Word and try running the script again.")
    finally:
        if os.path.exists(full_image_path):
            os.remove(full_image_path)


# Execution block handling your exact directory layout and custom string target:
insert_code128_barcode(
    'C:\\Users\\joseph.adogeri\\Desktop\\test\\pp_test.docx',
    'C:\\Users\\joseph.adogeri\\Desktop\\test\\final_notice.docx',
    '1234-JOSEPHADOGERIKINGST'
)
